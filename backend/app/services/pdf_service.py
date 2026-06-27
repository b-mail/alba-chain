"""PDF 처리: 텍스트 추출(pdfplumber, +vision 폴백)과 표준계약서 PDF 생성(WeasyPrint)."""
from __future__ import annotations

import logging
import uuid
from pathlib import Path
from typing import Any

from jinja2 import Environment, select_autoescape

from app.core.config import settings

logger = logging.getLogger(__name__)


def extract_text(path: Path) -> str:
    """업로드 문서에서 텍스트를 추출한다. 확장자(.pdf/.docx/.txt)로 분기."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix in (".txt", ".md"):
        try:
            return path.read_text(encoding="utf-8").strip()
        except Exception as exc:  # noqa: BLE001
            logger.warning("텍스트 파일 읽기 실패: %s", exc)
            return ""
    return extract_text_from_pdf(path)


def extract_text_from_pdf(pdf_path: Path) -> str:
    """텍스트 PDF 에서 본문을 추출한다. 스캔본 등 텍스트가 비면 빈 문자열."""
    try:
        import pdfplumber

        chunks: list[str] = []
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")
        text = "\n".join(chunks).strip()
        return text
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdfplumber 추출 실패: %s", exc)
        return ""


def extract_text_from_docx(docx_path: Path) -> str:
    """DOCX 에서 본문 단락 + 표 셀 텍스트를 추출한다(표준근로계약서는 표 기반)."""
    try:
        import docx

        document = docx.Document(str(docx_path))
        lines: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines).strip()
    except Exception as exc:  # noqa: BLE001
        logger.warning("docx 추출 실패: %s", exc)
        return ""


_CONTRACT_TEMPLATE = """<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="utf-8" />
<style>
  @page { size: A4; margin: 2cm; }
  body { font-family: "Noto Sans CJK KR", "AppleGothic", "Malgun Gothic", sans-serif; font-size: 11pt; color: #111; }
  h1 { text-align: center; font-size: 18pt; margin-bottom: 24px; }
  table { width: 100%; border-collapse: collapse; margin-top: 12px; }
  th, td { border: 1px solid #555; padding: 8px 10px; text-align: left; vertical-align: top; }
  th { width: 28%; background: #f2f2f2; }
  .sign { margin-top: 40px; }
  .sign-row { display: flex; justify-content: space-between; margin-top: 16px; }
</style>
</head>
<body>
  <h1>표준근로계약서</h1>
  <table>
    <tr><th>사업체명(갑)</th><td>{{ gap_company_name or '' }}</td></tr>
    <tr><th>근로자(을)</th><td>{{ eul_name or '' }}</td></tr>
    <tr><th>근로계약기간</th><td>{{ contract_start_date or '' }} ~ {{ contract_end_date or '기간의 정함 없음' }}</td></tr>
    <tr><th>근무장소</th><td>{{ workplace or '' }}</td></tr>
    <tr><th>업무내용</th><td>{{ job_description or '' }}</td></tr>
    <tr><th>근무일/근로시간</th><td>주 {{ work_days_per_week or '' }}일 ({{ work_days_note or '' }}), {{ work_start_time or '' }} ~ {{ work_end_time or '' }}</td></tr>
    <tr><th>휴게시간</th><td>{{ break_start_time or '-' }} ~ {{ break_end_time or '-' }}</td></tr>
    <tr><th>주휴일</th><td>{{ weekly_holiday_day or '' }}</td></tr>
    <tr><th>임금</th><td>{{ wage_type or '' }} {{ '{:,}'.format(wage_amount) if wage_amount else '' }} 원</td></tr>
    <tr><th>임금지급일</th><td>{{ pay_cycle or '' }} {{ pay_day or '' }}일, {{ pay_method or '' }}</td></tr>
    <tr><th>사회보험</th><td>
      고용보험 {{ '가입' if insurance_employment else '미가입' }} /
      산재보험 {{ '가입' if insurance_industrial else '미가입' }} /
      국민연금 {{ '가입' if insurance_pension else '미가입' }} /
      건강보험 {{ '가입' if insurance_health else '미가입' }}
    </td></tr>
  </table>
  <div class="sign">
    <p>위와 같이 근로계약을 체결한다.</p>
    <div class="sign-row"><span>작성일: {{ contract_date or '' }}</span></div>
    <div class="sign-row"><span>(갑) {{ gap_business_name or gap_company_name or '' }} (서명)</span><span>(을) {{ eul_full_name or eul_name or '' }} (서명)</span></div>
  </div>
</body>
</html>
"""


def generate_contract_pdf(terms: dict[str, Any], contract_id: int) -> str:
    """추출/확정 조건으로 표준계약서 PDF 를 생성하고 file_url 을 반환한다."""
    env = Environment(autoescape=select_autoescape(["html"]))
    template = env.from_string(_CONTRACT_TEMPLATE)
    html = template.render(**terms)

    filename = f"contract_{contract_id}_{uuid.uuid4().hex[:8]}.pdf"
    out_dir = settings.storage_path / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename

    try:
        from weasyprint import HTML

        HTML(string=html).write_pdf(str(out_path))
    except Exception as exc:  # noqa: BLE001  weasyprint 시스템 의존성 미설치 폴백
        logger.warning("WeasyPrint PDF 생성 실패, HTML 폴백 저장: %s", exc)
        out_path = out_path.with_suffix(".html")
        out_path.write_text(html, encoding="utf-8")

    return f"{settings.files_url_prefix}/generated/{out_path.name}"


def resolve_file_url(file_url: str) -> Path:
    """`/files/...` 형태의 file_url 을 로컬 절대경로로 변환한다."""
    prefix = settings.files_url_prefix.rstrip("/")
    rel = file_url
    if rel.startswith(prefix):
        rel = rel[len(prefix) :]
    rel = rel.lstrip("/")
    return settings.storage_path / rel


def save_upload(file_bytes: bytes, original_name: str, subdir: str = "uploads") -> tuple[Path, str]:
    """업로드 파일을 저장하고 (절대경로, file_url) 반환."""
    suffix = Path(original_name).suffix or ".bin"
    filename = f"{uuid.uuid4().hex}{suffix}"
    out_dir = settings.storage_path / subdir
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    out_path.write_bytes(file_bytes)
    return out_path, f"{settings.files_url_prefix}/{subdir}/{filename}"
